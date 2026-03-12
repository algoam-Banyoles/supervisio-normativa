"""
report.py — Generador d'informe Word (.docx) amb python-docx
"""

import os
import sys
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from pathlib import Path
from datetime import datetime
import re

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


C_OK    = RGBColor(0x1a, 0x7f, 0x4b)
C_NOOK  = RGBColor(0xc0, 0x39, 0x2b)
C_INFO  = RGBColor(0x1a, 0x5f, 0xa8)
C_TITLE = RGBColor(0x1e, 0x2d, 0x40)
C_WHITE = RGBColor(0xff, 0xff, 0xff)

STATUS_LABELS = {
    "OK": "✓ OK",
    "NO OK": "✗ NO OK",
    "INFO": "● INFO",
}


def _is_valid_xml_char(codepoint: int) -> bool:
    return (
        codepoint == 0x9
        or codepoint == 0xA
        or codepoint == 0xD
        or 0x20 <= codepoint <= 0xD7FF
        or 0xE000 <= codepoint <= 0xFFFD
        or 0x10000 <= codepoint <= 0x10FFFF
    )


def _xml_safe_text(value) -> str:
    if value is None:
        return ""
    text = str(value)
    return "".join(ch for ch in text if _is_valid_xml_char(ord(ch)))


class ReportGenerator:
    def __init__(self, filename: str, results: list[dict]):
        self.filename = filename
        self.results  = results
        self.ts       = datetime.now().strftime("%d/%m/%Y %H:%M")

    def save_docx(self, output_path: Path):
        doc = Document()
        _set_page_margins(doc, top=2.5, bottom=2.5, left=3.0, right=2.5)
        _set_default_font(doc)

        page_count = _guess_page_count(self.results)
        total = sum(len(r["findings"]) for r in self.results)
        nook  = sum(len([f for f in r["findings"] if f["status"] == "NO OK"]) for r in self.results)
        ok    = sum(len([f for f in r["findings"] if f["status"] == "OK"])    for r in self.results)
        info  = sum(len([f for f in r["findings"] if f["status"] == "INFO"])  for r in self.results)

        for section in doc.sections:
            _configure_header(section, self.filename)
            _configure_footer(section)

        _add_title_block(doc, self.filename, self.ts, page_count)
        _add_summary_table(doc, total, ok, nook, info)
        doc.add_paragraph()

        for section in self.results:
            _add_section(doc, section)
            doc.add_paragraph()

        _add_normativa_aplicada(doc, self.results)
        _add_most_relevant(doc, self.results)

        doc.save(str(output_path))


def _set_page_margins(doc, top, bottom, left, right):
    for section in doc.sections:
        section.top_margin    = Cm(top)
        section.bottom_margin = Cm(bottom)
        section.left_margin   = Cm(left)
        section.right_margin  = Cm(right)
        section.header_distance = Cm(1.0)
        section.footer_distance = Cm(1.0)


def _set_default_font(doc):
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(9)


def _configure_header(section, filename: str):
    header = section.header
    p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    p.clear()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)

    usable_width = section.page_width - section.left_margin - section.right_margin
    p.paragraph_format.tab_stops.add_tab_stop(usable_width, WD_TAB_ALIGNMENT.RIGHT)

    left = p.add_run(_xml_safe_text(filename))
    left.font.size = Pt(8)
    left.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    p.add_run("\t")

    right_lbl = p.add_run("Pagina ")
    right_lbl.font.size = Pt(8)
    right_lbl.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    _append_page_number_field(p)

    _add_bottom_border(p, color="B8BDC7", size=4)


def _configure_footer(section):
    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.clear()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Project Checker v1.0 — Informe orientatiu")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    run.font.italic = True


def _append_page_number_field(paragraph):
    run = paragraph.add_run()
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "

    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")

    txt = OxmlElement("w:t")
    txt.text = "1"

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(txt)
    run._r.append(fld_end)


def _add_title_block(doc, filename, ts, page_count):
    # Bloc superior: placeholder de logo + titol a la dreta
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False

    logo_cell = table.cell(0, 0)
    title_cell = table.cell(0, 1)
    logo_cell.width = Cm(4.0)
    title_cell.width = Cm(11.5)

    _set_cell_bg(logo_cell, "D9DDE3")
    logo_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p_logo = logo_cell.paragraphs[0]
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_logo = p_logo.add_run("LOGO")
    r_logo.font.bold = True
    r_logo.font.size = Pt(11)
    r_logo.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    title_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p_title = title_cell.paragraphs[0]
    p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r_title = p_title.add_run("INFORME DE SUPERVISIO TECNICA DEL PROJECTE")
    r_title.font.bold = True
    r_title.font.size = Pt(16)
    r_title.font.color.rgb = C_TITLE

    p_name = title_cell.add_paragraph()
    p_name.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r_name = p_name.add_run(_xml_safe_text(filename))
    r_name.font.bold = True
    r_name.font.size = Pt(11)
    r_name.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    p_meta = doc.add_paragraph()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r_meta = p_meta.add_run(
        f"Data: {ts}    |    Fitxer: {_xml_safe_text(filename)}    |    Pagines analitzades: {page_count}"
    )
    r_meta.font.size = Pt(8)
    r_meta.font.color.rgb = RGBColor(0x7a, 0x7a, 0x7a)
    _add_bottom_border(p_meta, color="1E2D40", size=16)

    doc.add_paragraph()


def _add_summary_table(doc, total, ok, nook, info):
    p = doc.add_paragraph()
    run = p.add_run("RESUM EXECUTIU")
    run.font.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = C_TITLE

    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False

    labels = ["TOTAL", "OK", "NO OK", "INFO"]
    values = [total, ok, nook, info]
    fills = ["E4E7EC", "E8F5EE", "FDF0EF", "EAF2FB"]
    txt_colors = [RGBColor(0x33, 0x33, 0x33), C_OK, C_NOOK, C_INFO]
    widths = [Cm(3.7), Cm(3.7), Cm(3.7), Cm(3.7)]

    for i, cell in enumerate(table.rows[0].cells):
        cell.width = widths[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _set_cell_bg(cell, fills[i])

        p_num = cell.paragraphs[0]
        p_num.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_num = p_num.add_run(str(values[i]))
        r_num.font.size = Pt(20)
        r_num.font.bold = True
        r_num.font.color.rgb = txt_colors[i]

        p_lbl = cell.add_paragraph()
        p_lbl.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_lbl = p_lbl.add_run(labels[i])
        r_lbl.font.size = Pt(8)
        r_lbl.font.bold = True
        r_lbl.font.color.rgb = txt_colors[i]


def _add_section(doc, section: dict):
    _add_section_title(doc, section["title"])

    findings = section["findings"]
    if not findings:
        return

    table  = doc.add_table(rows=1 + len(findings), cols=5)
    table.style = "Table Grid"
    table.autofit = False
    widths = [Cm(2.0), Cm(1.8), Cm(6.0), Cm(5.0), Cm(3.5)]

    for i, cell in enumerate(table.rows[0].cells):
        cell.width = widths[i]
        _set_cell_bg(cell, "E9ECF1")
        run = cell.paragraphs[0].add_run(["Estat", "Item", "Descripcio", "Detall", "Referencia"][i])
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run.font.bold = True
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x44,0x44,0x44)

    for row_idx, finding in enumerate(findings, 1):
        row    = table.rows[row_idx]
        cells  = row.cells
        status = finding.get("status", "INFO")
        bg     = "FFFFFF" if row_idx % 2 == 1 else "F7F8FA"

        for cell in cells:
            _set_cell_bg(cell, bg)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

        if status == "NO OK":
            _set_cell_left_border(cells[0], "C0392B", sz=16)

        # Estat
        cells[0].width = widths[0]
        p0 = cells[0].paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r0 = p0.add_run(STATUS_LABELS.get(status, status))
        r0.font.size = Pt(9)
        r0.font.bold = True
        r0.font.color.rgb = {"OK":C_OK,"NO OK":C_NOOK,"INFO":C_INFO}.get(status, C_INFO)

        # Item
        cells[1].width = widths[1]
        cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = cells[1].paragraphs[0].add_run(_xml_safe_text(finding.get("item", "")))
        r1.font.size = Pt(8)
        r1.font.name = "Courier New"
        r1.font.color.rgb = RGBColor(0x77,0x77,0x77)

        # Descripcio
        cells[2].width = widths[2]
        r2 = cells[2].paragraphs[0].add_run(_xml_safe_text(finding.get("descrip", "")))
        r2.font.size = Pt(9)

        # Detall
        cells[3].width = widths[3]
        detall = _xml_safe_text(finding.get("detall", "")).strip()
        r3 = cells[3].paragraphs[0].add_run(detall)
        r3.font.size = Pt(8)
        r3.font.name = "Courier New"
        r3.font.color.rgb = RGBColor(0x55,0x55,0x55)

        # Referencia
        cells[4].width = widths[4]
        r4 = cells[4].paragraphs[0].add_run(_xml_safe_text(finding.get("ref", "")))
        r4.font.size   = Pt(8)
        r4.font.italic = True
        r4.font.color.rgb = RGBColor(0x88,0x88,0x88)


def _add_most_relevant(doc, results: list[dict]):
    _add_section_title(doc, "ASPECTES MES RELLEVANTS")

    for section in results:
        local_nook = [f for f in section.get("findings", []) if f.get("status") == "NO OK"]
        if not local_nook:
            continue

        p_sec = doc.add_paragraph()
        r_sec = p_sec.add_run("▸ " + _xml_safe_text(section.get("title", "")))
        r_sec.font.bold = True
        r_sec.font.size = Pt(9)
        r_sec.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
        p_sec.paragraph_format.space_before = Pt(4)
        p_sec.paragraph_format.space_after = Pt(2)

        for finding in local_nook:
            item = _xml_safe_text(finding.get("item", ""))
            descrip = _xml_safe_text(finding.get("descrip", ""))
            bullet = doc.add_paragraph(style="List Bullet")
            rb = bullet.add_run(f"[{item}] {descrip}")
            rb.font.size = Pt(9)


def _add_normativa_aplicada(doc, results: list[dict]):
    rows = _extract_normativa_rows(results)
    _add_section_title(doc, "TAULA DE NORMATIVA APLICADA")

    if not rows:
        p = doc.add_paragraph()
        run = p.add_run("No s'han pogut obtenir referencies normatives estructurades per a aquesta execucio.")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        _add_normativa_note(doc)
        return

    table = doc.add_table(rows=1 + len(rows), cols=5)
    table.style = "Table Grid"
    table.autofit = False
    widths = [Cm(5.0), Cm(2.0), Cm(2.5), Cm(2.5), Cm(5.5)]
    headers = [
        "Referencia normativa",
        "Tipus",
        "Estat",
        "Pagines on apareix",
        "Observacions",
    ]

    for i, cell in enumerate(table.rows[0].cells):
        cell.width = widths[i]
        _set_cell_bg(cell, "E9ECF1")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(headers[i])
        r.font.size = Pt(8)
        r.font.bold = True
        r.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    for idx, row in enumerate(rows, 1):
        cells = table.rows[idx].cells
        for i, cell in enumerate(cells):
            cell.width = widths[i]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

        status = row.get("status", "PENDENT")
        if status == "DEROGADA":
            bg = "FDF0EF"
            status_color = C_NOOK
        elif status == "PENDENT":
            bg = "FFFDE7"
            status_color = RGBColor(0xC2, 0x7C, 0x00)
        else:
            bg = "FFFFFF"
            status_color = C_OK

        for cell in cells:
            _set_cell_bg(cell, bg)

        _add_table_text(cells[0], row.get("reference", ""), 8.5)
        _add_table_text(cells[1], row.get("type", "Altres"), 8.5)
        _add_table_text(cells[2], status, 8.5, bold=True, color=status_color)
        _add_table_text(cells[3], _format_normativa_pages(row.get("pages", [])), 8.0)
        _add_table_text(
            cells[4],
            row.get("observations", ""),
            8.0,
            bold=(status == "DEROGADA"),
            color=(status_color if status == "DEROGADA" else RGBColor(0x55, 0x55, 0x55)),
        )

    _add_normativa_note(doc)


def _extract_normativa_rows(results: list[dict]) -> list[dict]:
    for section in results:
        for finding in section.get("findings", []):
            rows = finding.get("normativa_rows")
            if isinstance(rows, list):
                return rows
    return []


def _add_table_text(cell, text: str, size: float, bold: bool = False, color=None):
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(_xml_safe_text(text))
    run.font.size = Pt(size)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color


def _format_normativa_pages(pages: list[int]) -> str:
    return ", ".join(str(p) for p in pages[:20]) if pages else "-"


def _add_normativa_note(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(
        "Taula generada automàticament. PENDENT DE VERIFICACIÓ indica normes no presents al "
        "catàleg intern (normativa_annexes.json). Verificació final a càrrec de l'enginyer revisor."
    )
    run.font.size = Pt(8)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)


def _add_section_title(doc, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    _set_para_bg(p, "1E2D40")

    run = p.add_run(text)
    run.font.name  = "Calibri"
    run.font.size  = Pt(10)
    run.font.bold  = True
    run.font.color.rgb = C_WHITE


def _guess_page_count(results: list[dict]) -> int:
    max_page = 0
    page_re = re.compile(r"\b(?:pag(?:ina|ines)?|p\.)\s*([0-9]{1,5})", re.IGNORECASE)

    for section in results:
        for finding in section.get("findings", []):
            text = f"{finding.get('descrip', '')}\n{finding.get('detall', '')}"
            for m in page_re.finditer(text):
                try:
                    max_page = max(max_page, int(m.group(1)))
                except ValueError:
                    continue

    return max_page if max_page > 0 else 0


def _set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def _set_cell_left_border(cell, color="C0392B", sz=16):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)

    left = tcBorders.find(qn("w:left"))
    if left is None:
        left = OxmlElement("w:left")
        tcBorders.append(left)

    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(sz))
    left.set(qn("w:space"), "0")
    left.set(qn("w:color"), color)


def _set_para_bg(para, hex_color: str):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    pPr.append(shd)


def _add_bottom_border(para, color="2E75B6", size=6):
    pPr    = para._p.get_or_add_pPr()
    pBdr   = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    str(size))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)
